import csv
from collections import defaultdict
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches


def select_csv_file():
    csv_file_path.set(filedialog.askopenfilename(filetypes=[("CSV Files", "*.csv")]))

def select_word_file():
    word_file_path.set(filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")]))

def select_output_file():
    output_file_path.set(filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")]))

def group_csv_data():
    grouped_data = defaultdict(list)
    with open(csv_file_path.get(), newline='', encoding='utf-8-sig') as csvfile:  # Use utf-8-sig to handle BOM
        reader = csv.DictReader(csvfile)
        for row in reader:
            # Normalize the column names by stripping white spaces and BOM
            normalized_row = {k.strip('\ufeff').strip(): v for k, v in row.items()}
            grouped_data[normalized_row['Module']].append(normalized_row)
    return grouped_data


def set_cell_background(cell, color_str):
    """Set background color for a cell."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_str))
    cell._tc.get_or_add_tcPr().append(shading_elm)


def replace_text_in_paragraph(paragraph, search_text, replace_text):
    for run in paragraph.runs:
        if search_text in run.text:
            run.text = run.text.replace(search_text, replace_text)


def insert_data_to_word(grouped_data, completion_date, customer_name):
    document = Document(word_file_path.get())

    # Find and replace "Epic name" with "OPP",and "Month-Year" in the entire document
    for paragraph in document.paragraphs:
        replace_text_in_paragraph(paragraph, "Completion date", completion_date.get())
        replace_text_in_paragraph(paragraph, "Epic name", customer_name.get())

    for module in grouped_data.keys():
        document.add_paragraph(f'{module}', style='List Number1')
    section_number = 1  # Starting section number
    for module, rows in grouped_data.items():
        # Add a heading for each module with an indent
        heading_text = f"{module}"
        document.add_heading(heading_text, level=2)
        #heading_paragraph = document.add_paragraph(heading_text, style=None)
        #heading_paragraph.paragraph_format.left_indent = Pt(20)

        # Add the specified line before the table
        line_paragraph = document.add_paragraph("The data migration included creation of the following objects:", style=None)
        line_paragraph.paragraph_format.left_indent = Pt(70)

        # Add the table
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'  # This applies a basic grid to the table
        hdr_cells = table.rows[0].cells

        # Set the table column widths
        for i, width in enumerate([Inches(2), Inches(1.5), Inches(3)]):
            table.columns[i].width = width

        # Style header cells
        for cell in hdr_cells:
            cell.text = 'Header'
            set_cell_background(cell, 'D9D9D9')  # Light grey background

        hdr_cells[0].text = 'Object'
        hdr_cells[1].text = 'Loaded'
        hdr_cells[2].text = 'Commentary'

        # Fill in table rows
        for row in rows:
            row_cells = table.add_row().cells
            row_cells[0].text = row.get('Object', '')
            row_cells[1].text = str(row.get('Loaded', ''))
            row_cells[2].text = row.get('Commentary', '') or ''

        # Add the upload comments section after each table
        comment_heading = document.add_paragraph('Upload comments:', style=None)
        comment_heading.paragraph_format.left_indent = Pt(70)
        comment_paragraph = document.add_paragraph('- N/A', style=None)
        comment_paragraph.paragraph_format.left_indent = Pt(75)

        # Add an empty line for separation
        document.add_paragraph()

        section_number += 1  # Increment section number for next module

    # Add the "Uploading verification" section at the end of the document
    # document.add_page_break()  # Add a page break before the new section if needed
    document.add_heading('Uploading verification:', level=1)
    verification_paragraph = document.add_paragraph()
    verification_paragraph.add_run('- DOT performed verification tests for a few random records.\n')
    verification_paragraph.add_run('- It is ')
    verification_paragraph.add_run('customer’s').bold = True
    verification_paragraph.add_run(' responsibility to make full verification to the uploaded data if needed, ')
    verification_paragraph.add_run('according to the customer’s procedures.').italic = True

    document.save(output_file_path.get())
    messagebox.showinfo("Success", "CSV data inserted into WORD document")


def combine_files():
    if not csv_file_path.get() or not word_file_path.get() or not output_file_path.get():
        messagebox.showerror("Error", "Please select all files")
        return

    grouped_data = group_csv_data()
    if grouped_data is not None:
        insert_data_to_word(grouped_data, completion_date, customer_name)

app = tk.Tk()
app.title("WORD and CSV Combiner")

csv_file_path = tk.StringVar()
word_file_path = tk.StringVar()
output_file_path = tk.StringVar()

tk.Label(app, text="CSV File Path:").pack()
tk.Entry(app, textvariable=csv_file_path).pack()
tk.Button(app, text="Browse", command=select_csv_file).pack()

tk.Label(app, text="WORD File Path:").pack()
tk.Entry(app, textvariable=word_file_path).pack()
tk.Button(app, text="Browse", command=select_word_file).pack()

tk.Label(app, text="Output File Path:").pack()
tk.Entry(app, textvariable=output_file_path).pack()
tk.Button(app, text="Browse", command=select_output_file).pack()

# Customer name input
tk.Label(app, text="Customer Name:").pack()
customer_name = tk.StringVar()
customer_name_entry = tk.Entry(app, textvariable=customer_name)
customer_name_entry.pack()

# Completion date input
tk.Label(app, text="Completion Date:").pack()
completion_date = tk.StringVar()
completion_date_entry = tk.Entry(app, textvariable=completion_date)
completion_date_entry.pack()


tk.Button(app, text="Combine Files", command=combine_files).pack()

app.mainloop()
